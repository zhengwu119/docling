import logging
import re
from copy import deepcopy
from io import BytesIO
from pathlib import Path
from typing import Any, Callable, Final, Optional, Union

from docling_core.types.doc import (
    ContentLayer,
    DocItemLabel,
    DoclingDocument,
    DocumentOrigin,
    GroupLabel,
    ImageRef,
    ListGroup,
    NodeItem,
    RefItem,
    RichTableCell,
    TableCell,
    TableData,
    TableItem,
)
from docling_core.types.doc.document import Formatting, Script
from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.table import CT_Tc
from docx.oxml.xmlchemy import BaseOxmlElement
from docx.table import Table, _Cell
from docx.text.hyperlink import Hyperlink
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from lxml import etree
from PIL import Image, UnidentifiedImageError
from pydantic import AnyUrl
from typing_extensions import override

from docling.backend.abstract_backend import DeclarativeDocumentBackend
from docling.backend.docx.drawingml.utils import (
    get_docx_to_pdf_converter,
    get_pil_from_dml_docx,
)
from docling.backend.docx.latex.omml import oMath2Latex
from docling.datamodel.base_models import InputFormat
from docling.datamodel.document import InputDocument

_log = logging.getLogger(__name__)


class MsWordDocumentBackend(DeclarativeDocumentBackend):
    _BLIP_NAMESPACES: Final = {
        "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
        "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
        "mc": "http://schemas.openxmlformats.org/markup-compatibility/2006",
        "v": "urn:schemas-microsoft-com:vml",
        "wps": "http://schemas.microsoft.com/office/word/2010/wordprocessingShape",
        "w10": "urn:schemas-microsoft-com:office:word",
        "a14": "http://schemas.microsoft.com/office/drawing/2010/main",
    }

    @override
    def __init__(
        self, in_doc: "InputDocument", path_or_stream: Union[BytesIO, Path]
    ) -> None:
        super().__init__(in_doc, path_or_stream)
        self.XML_KEY = (
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val"
        )
        self.xml_namespaces = {
            "w": "http://schemas.microsoft.com/office/word/2003/wordml"
        }
        self.blip_xpath_expr = etree.XPath(
            ".//a:blip", namespaces=MsWordDocumentBackend._BLIP_NAMESPACES
        )
        # self.initialise(path_or_stream)
        # Word file:
        self.path_or_stream: Union[BytesIO, Path] = path_or_stream
        self.valid: bool = False
        # Initialise the parents for the hierarchy
        self.max_levels: int = 10
        self.level_at_new_list: Optional[int] = None
        self.parents: dict[int, Optional[NodeItem]] = {}
        self.numbered_headers: dict[int, int] = {}
        self.equation_bookends: str = "<eq>{EQ}</eq>"
        # Track processed textbox elements to avoid duplication
        self.processed_textbox_elements: list[int] = []
        self.docx_to_pdf_converter: Optional[Callable] = None
        self.docx_to_pdf_converter_init = False
        self.display_drawingml_warning = True

        for i in range(-1, self.max_levels):
            self.parents[i] = None

        self.level = 0
        self.listIter = 0
        # Track list counters per numId and ilvl
        self.list_counters: dict[tuple[int, int], int] = {}
        # Set starting content layer
        self.content_layer = ContentLayer.BODY

        self.history: dict[str, Any] = {
            "names": [None],
            "levels": [None],
            "numids": [None],
            "indents": [None],
        }

        self.docx_obj = self.load_msword_file(
            path_or_stream=self.path_or_stream, document_hash=self.document_hash
        )
        if self.docx_obj:
            self.valid = True

    @override
    def is_valid(self) -> bool:
        return self.valid

    @classmethod
    @override
    def supports_pagination(cls) -> bool:
        return False

    @override
    def unload(self):
        if isinstance(self.path_or_stream, BytesIO):
            self.path_or_stream.close()

        self.path_or_stream = None

    @classmethod
    @override
    def supported_formats(cls) -> set[InputFormat]:
        return {InputFormat.DOCX}

    @override
    def convert(self) -> DoclingDocument:
        """Parses the DOCX into a structured document model.

        Returns:
            The parsed document.
        """

        origin = DocumentOrigin(
            filename=self.file.name or "file",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            binary_hash=self.document_hash,
        )

        doc = DoclingDocument(name=self.file.stem or "file", origin=origin)
        if self.is_valid():
            assert self.docx_obj is not None
            doc, _ = self._walk_linear(self.docx_obj.element.body, doc)
            self._add_header_footer(self.docx_obj, doc)

            return doc
        else:
            raise RuntimeError(
                f"Cannot convert doc with {self.document_hash} because the backend failed to init."
            )

    @staticmethod
    def load_msword_file(
        path_or_stream: Union[BytesIO, Path], document_hash: str
    ) -> DocxDocument:
        try:
            if isinstance(path_or_stream, BytesIO):
                return Document(path_or_stream)
            elif isinstance(path_or_stream, Path):
                return Document(str(path_or_stream))
            else:
                return None
        except Exception as e:
            raise RuntimeError(
                f"MsWordDocumentBackend could not load document with hash {document_hash}"
            ) from e

    def _update_history(
        self,
        name: str,
        level: Optional[int],
        numid: Optional[int],
        ilevel: Optional[int],
    ):
        self.history["names"].append(name)
        self.history["levels"].append(level)

        self.history["numids"].append(numid)
        self.history["indents"].append(ilevel)

    def _prev_name(self) -> Optional[str]:
        return self.history["names"][-1]

    def _prev_level(self) -> Optional[int]:
        return self.history["levels"][-1]

    def _prev_numid(self) -> Optional[int]:
        return self.history["numids"][-1]

    def _prev_indent(self) -> Optional[int]:
        return self.history["indents"][-1]

    def _get_level(self) -> int:
        """Return the first None index."""
        for k, v in self.parents.items():
            if k >= 0 and v is None:
                return k
        return 0

    def _walk_linear(
        self,
        body: BaseOxmlElement,
        doc: DoclingDocument,
        # parent:
    ) -> tuple[DoclingDocument, list[RefItem]]:
        added_elements = []
        for element in body:
            tag_name = etree.QName(element).localname
            # Check for Inline Images (blip elements)
            drawing_blip = self.blip_xpath_expr(element)
            drawingml_els = element.findall(
                ".//w:drawing", namespaces=MsWordDocumentBackend._BLIP_NAMESPACES
            )

            # Check for textbox content - check multiple textbox formats
            # Only process if the element hasn't been processed before
            element_id = id(element)
            if element_id not in self.processed_textbox_elements:
                # Modern Word textboxes
                txbx_xpath = etree.XPath(
                    ".//w:txbxContent|.//v:textbox//w:p",
                    namespaces=MsWordDocumentBackend._BLIP_NAMESPACES,
                )
                textbox_elements = txbx_xpath(element)

                # No modern textboxes found, check for alternate/legacy textbox formats
                if not textbox_elements and tag_name in ["drawing", "pict"]:
                    # Additional checks for textboxes in DrawingML and VML formats
                    alt_txbx_xpath = etree.XPath(
                        ".//wps:txbx//w:p|.//w10:wrap//w:p|.//a:p//a:t",
                        namespaces=MsWordDocumentBackend._BLIP_NAMESPACES,
                    )
                    textbox_elements = alt_txbx_xpath(element)

                    # Check for shape text that's not in a standard textbox
                    if not textbox_elements:
                        shape_text_xpath = etree.XPath(
                            ".//a:bodyPr/ancestor::*//a:t|.//a:txBody//a:t",
                            namespaces=MsWordDocumentBackend._BLIP_NAMESPACES,
                        )
                        shape_text_elements = shape_text_xpath(element)
                        if shape_text_elements:
                            # Create custom text elements from shape text
                            text_content = " ".join(
                                [t.text for t in shape_text_elements if t.text]
                            )
                            if text_content.strip():
                                _log.debug(f"Found shape text: {text_content[:50]}...")
                                # Create a paragraph-like element to process with standard handler
                                level = self._get_level()
                                shape_group = doc.add_group(
                                    label=GroupLabel.SECTION,
                                    parent=self.parents[level - 1],
                                    name="shape-text",
                                    content_layer=self.content_layer,
                                )
                                added_elements.append(shape_group.get_ref())
                                doc.add_text(
                                    label=DocItemLabel.TEXT,
                                    parent=shape_group,
                                    text=text_content,
                                    content_layer=self.content_layer,
                                )

                if textbox_elements:
                    # Mark the parent element as processed
                    self.processed_textbox_elements.append(element_id)
                    # Also mark all found textbox elements as processed
                    for tb_element in textbox_elements:
                        self.processed_textbox_elements.append(id(tb_element))

                    _log.debug(
                        f"Found textbox content with {len(textbox_elements)} elements"
                    )
                    tbc = self._handle_textbox_content(textbox_elements, doc)
                    added_elements.extend(tbc)

            # Check for Tables
            if tag_name == "tbl":
                try:
                    t = self._handle_tables(element, doc)
                    added_elements.extend(t)
                except Exception:
                    _log.debug("could not parse a table, broken docx table")
            # Check for Image
            elif drawing_blip:
                pics = self._handle_pictures(drawing_blip, doc)
                added_elements.extend(pics)
                # Check for Text after the Image
                if (
                    tag_name == "p"
                    and element.find(
                        ".//w:t", namespaces=MsWordDocumentBackend._BLIP_NAMESPACES
                    )
                    is not None
                ):
                    te1 = self._handle_text_elements(element, doc)
                    added_elements.extend(te1)
            # Check for DrawingML elements
            elif drawingml_els:
                if (
                    self.docx_to_pdf_converter is None
                    and self.docx_to_pdf_converter_init is False
                ):
                    self.docx_to_pdf_converter = get_docx_to_pdf_converter()
                    self.docx_to_pdf_converter_init = True

                if self.docx_to_pdf_converter is None:
                    if self.display_drawingml_warning:
                        if self.docx_to_pdf_converter is None:
                            _log.warning(
                                "Found DrawingML elements in document, but no DOCX to PDF converters. "
                                "If you want these exported, make sure you have "
                                "LibreOffice binary in PATH or specify its path with DOCLING_LIBREOFFICE_CMD."
                            )
                            self.display_drawingml_warning = False
                else:
                    self._handle_drawingml(doc=doc, drawingml_els=drawingml_els)
            # Check for the sdt containers, like table of contents
            elif tag_name == "sdt":
                sdt_content = element.find(
                    ".//w:sdtContent", namespaces=MsWordDocumentBackend._BLIP_NAMESPACES
                )
                if sdt_content is not None:
                    # Iterate paragraphs, runs, or text inside <w:sdtContent>.
                    paragraphs = sdt_content.findall(
                        ".//w:p", namespaces=MsWordDocumentBackend._BLIP_NAMESPACES
                    )
                    for p in paragraphs:
                        te = self._handle_text_elements(p, doc)
                        added_elements.extend(te)
            # Check for Text
            elif tag_name == "p":
                # "tcPr", "sectPr"
                te = self._handle_text_elements(element, doc)
                added_elements.extend(te)
            else:
                _log.debug(f"Ignoring element in DOCX with tag: {tag_name}")

        return doc, added_elements

    def _str_to_int(
        self, s: Optional[str], default: Optional[int] = 0
    ) -> Optional[int]:
        if s is None:
            return None
        try:
            return int(s)
        except ValueError:
            return default

    def _split_text_and_number(self, input_string: str) -> list[str]:
        match = re.match(r"(\D+)(\d+)$|^(\d+)(\D+)", input_string)
        if match:
            parts = list(filter(None, match.groups()))
            return parts
        else:
            return [input_string]

    def _get_numId_and_ilvl(
        self, paragraph: Paragraph
    ) -> tuple[Optional[int], Optional[int]]:
        # Access the XML element of the paragraph
        numPr = paragraph._element.find(
            ".//w:numPr", namespaces=paragraph._element.nsmap
        )

        if numPr is not None:
            # Get the numId element and extract the value
            numId_elem = numPr.find("w:numId", namespaces=paragraph._element.nsmap)
            ilvl_elem = numPr.find("w:ilvl", namespaces=paragraph._element.nsmap)
            numId = numId_elem.get(self.XML_KEY) if numId_elem is not None else None
            ilvl = ilvl_elem.get(self.XML_KEY) if ilvl_elem is not None else None

            return self._str_to_int(numId, None), self._str_to_int(ilvl, None)

        return None, None  # If the paragraph is not part of a list

    def _get_list_counter(self, numid: int, ilvl: int) -> int:
        """Get and increment the counter for a specific numId and ilvl combination."""
        key = (numid, ilvl)
        if key not in self.list_counters:
            self.list_counters[key] = 0
        self.list_counters[key] += 1
        return self.list_counters[key]

    def _reset_list_counters_for_new_sequence(self, numid: int):
        """Reset counters when starting a new numbering sequence."""
        # Reset all counters for this numid
        keys_to_reset = [key for key in self.list_counters.keys() if key[0] == numid]
        for key in keys_to_reset:
            self.list_counters[key] = 0

    def _is_numbered_list(self, numId: int, ilvl: int) -> bool:
        """Check if a list is numbered based on its numFmt value."""
        try:
            # Access the numbering part of the document
            if not hasattr(self.docx_obj, "part") or not hasattr(
                self.docx_obj.part, "package"
            ):
                return False

            numbering_part = None
            # Find the numbering part
            for part in self.docx_obj.part.package.parts:
                if "numbering" in part.partname:
                    numbering_part = part
                    break

            if numbering_part is None:
                return False

            # Parse the numbering XML
            numbering_root = numbering_part.element
            namespaces = {
                "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
            }

            # Find the numbering definition with the given numId
            num_xpath = f".//w:num[@w:numId='{numId}']"
            num_element = numbering_root.find(num_xpath, namespaces=namespaces)

            if num_element is None:
                return False

            # Get the abstractNumId from the num element
            abstract_num_id_elem = num_element.find(
                ".//w:abstractNumId", namespaces=namespaces
            )
            if abstract_num_id_elem is None:
                return False

            abstract_num_id = abstract_num_id_elem.get(
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val"
            )
            if abstract_num_id is None:
                return False

            # Find the abstract numbering definition
            abstract_num_xpath = (
                f".//w:abstractNum[@w:abstractNumId='{abstract_num_id}']"
            )
            abstract_num_element = numbering_root.find(
                abstract_num_xpath, namespaces=namespaces
            )

            if abstract_num_element is None:
                return False

            # Find the level definition for the given ilvl
            lvl_xpath = f".//w:lvl[@w:ilvl='{ilvl}']"
            lvl_element = abstract_num_element.find(lvl_xpath, namespaces=namespaces)

            if lvl_element is None:
                return False

            # Get the numFmt element
            num_fmt_element = lvl_element.find(".//w:numFmt", namespaces=namespaces)
            if num_fmt_element is None:
                return False

            num_fmt = num_fmt_element.get(
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val"
            )

            # Numbered formats include: decimal, lowerRoman, upperRoman, lowerLetter, upperLetter
            # Bullet formats include: bullet
            numbered_formats = {
                "decimal",
                "lowerRoman",
                "upperRoman",
                "lowerLetter",
                "upperLetter",
                "decimalZero",
            }

            return num_fmt in numbered_formats

        except Exception as e:
            _log.debug(f"Error determining if list is numbered: {e}")
            return False

    def _get_heading_and_level(self, style_label: str) -> tuple[str, Optional[int]]:
        parts = self._split_text_and_number(style_label)

        if len(parts) == 2:
            parts.sort()
            label_str: str = ""
            label_level: Optional[int] = 0
            if parts[0].strip().lower() == "heading":
                label_str = "Heading"
                label_level = self._str_to_int(parts[1], None)
            if parts[1].strip().lower() == "heading":
                label_str = "Heading"
                label_level = self._str_to_int(parts[0], None)
            return label_str, label_level

        return style_label, None

    def _get_label_and_level(self, paragraph: Paragraph) -> tuple[str, Optional[int]]:
        if paragraph.style is None:
            return "Normal", None

        label = paragraph.style.style_id
        name = paragraph.style.name
        base_style_label = None
        base_style_name = None
        if base_style := getattr(paragraph.style, "base_style", None):
            base_style_label = base_style.style_id
            base_style_name = base_style.name

        if label is None:
            return "Normal", None

        if ":" in label:
            parts = label.split(":")
            if len(parts) == 2:
                return parts[0], self._str_to_int(parts[1], None)

        if "heading" in label.lower():
            return self._get_heading_and_level(label)
        if "heading" in name.lower():
            return self._get_heading_and_level(name)
        if base_style_label and "heading" in base_style_label.lower():
            return self._get_heading_and_level(base_style_label)
        if base_style_name and "heading" in base_style_name.lower():
            return self._get_heading_and_level(base_style_name)

        return label, None

    @classmethod
    def _get_format_from_run(cls, run: Run) -> Optional[Formatting]:
        # The .bold and .italic properties are booleans, but .underline can be an enum
        # like WD_UNDERLINE.THICK (value 6), so we need to convert it to a boolean
        is_bold = run.bold or False
        is_italic = run.italic or False
        is_strikethrough = run.font.strike or False
        # Convert any non-None underline value to True
        is_underline = bool(run.underline is not None and run.underline)
        is_sub = run.font.subscript or False
        is_sup = run.font.superscript or False
        script = Script.SUB if is_sub else Script.SUPER if is_sup else Script.BASELINE

        return Formatting(
            bold=is_bold,
            italic=is_italic,
            underline=is_underline,
            strikethrough=is_strikethrough,
            script=script,
        )

    def _get_paragraph_elements(self, paragraph: Paragraph):
        """
        Extract paragraph elements along with their formatting and hyperlink
        """

        # for now retain empty paragraphs for backwards compatibility:
        if paragraph.text.strip() == "":
            return [("", None, None)]

        paragraph_elements: list[
            tuple[str, Optional[Formatting], Optional[Union[AnyUrl, Path]]]
        ] = []
        group_text = ""
        previous_format = None

        # Iterate over the runs of the paragraph and group them by format
        for c in paragraph.iter_inner_content():
            if isinstance(c, Hyperlink):
                text = c.text
                hyperlink = Path(c.address)
                format = (
                    self._get_format_from_run(c.runs[0])
                    if c.runs and len(c.runs) > 0
                    else None
                )
            elif isinstance(c, Run):
                text = c.text
                hyperlink = None
                format = self._get_format_from_run(c)
            else:
                continue

            if (len(text.strip()) and format != previous_format) or (
                hyperlink is not None
            ):
                # If the style changes for a non empty text, add the previous group
                if len(group_text.strip()) > 0:
                    paragraph_elements.append(
                        (group_text.strip(), previous_format, None)
                    )
                group_text = ""

                # If there is a hyperlink, add it immediately
                if hyperlink is not None:
                    paragraph_elements.append((text.strip(), format, hyperlink))
                    text = ""
                else:
                    previous_format = format

            group_text += text

        # Format the last group
        if len(group_text.strip()) > 0:
            paragraph_elements.append((group_text.strip(), format, None))

        return paragraph_elements

    def _get_paragraph_position(self, paragraph_element):
        """Extract vertical position information from paragraph element."""
        # First try to directly get the index from w:p element that has an order-related attribute
        if (
            hasattr(paragraph_element, "getparent")
            and paragraph_element.getparent() is not None
        ):
            parent = paragraph_element.getparent()
            # Get all paragraph siblings
            paragraphs = [
                p for p in parent.getchildren() if etree.QName(p).localname == "p"
            ]
            # Find index of current paragraph within its siblings
            try:
                paragraph_index = paragraphs.index(paragraph_element)
                return paragraph_index  # Use index as position for consistent ordering
            except ValueError:
                pass

        # Look for position hints in element attributes and ancestor elements
        for elem in (*[paragraph_element], *paragraph_element.iterancestors()):
            # Check for direct position attributes
            for attr_name in ["y", "top", "positionY", "y-position", "position"]:
                value = elem.get(attr_name)
                if value:
                    try:
                        # Remove any non-numeric characters (like 'pt', 'px', etc.)
                        clean_value = re.sub(r"[^0-9.]", "", value)
                        if clean_value:
                            return float(clean_value)
                    except (ValueError, TypeError):
                        pass

            # Check for position in transform attribute
            transform = elem.get("transform")
            if transform:
                # Extract translation component from transform matrix
                match = re.search(r"translate\([^,]+,\s*([0-9.]+)", transform)
                if match:
                    try:
                        return float(match.group(1))
                    except ValueError:
                        pass

            # Check for anchors or relative position indicators in Word format
            # 'dist' attributes can indicate relative positioning
            for attr_name in ["distT", "distB", "anchor", "relativeFrom"]:
                if elem.get(attr_name) is not None:
                    return elem.sourceline  # Use the XML source line number as fallback

        # For VML shapes, look for specific attributes
        for ns_uri in paragraph_element.nsmap.values():
            if "vml" in ns_uri:
                # Try to extract position from style attribute
                style = paragraph_element.get("style")
                if style:
                    match = re.search(r"top:([0-9.]+)pt", style)
                    if match:
                        try:
                            return float(match.group(1))
                        except ValueError:
                            pass

        # If no better position indicator found, use XML source line number as proxy for order
        return (
            paragraph_element.sourceline
            if hasattr(paragraph_element, "sourceline")
            else None
        )

    def _collect_textbox_paragraphs(self, textbox_elements):
        """Collect and organize paragraphs from textbox elements."""
        processed_paragraphs = []
        container_paragraphs = {}

        for element in textbox_elements:
            element_id = id(element)
            # Skip if we've already processed this exact element
            if element_id in processed_paragraphs:
                continue

            tag_name = etree.QName(element).localname
            processed_paragraphs.append(element_id)

            # Handle paragraphs directly found (VML textboxes)
            if tag_name == "p":
                # Find the containing textbox or shape element
                container_id = None
                for ancestor in element.iterancestors():
                    if any(ns in ancestor.tag for ns in ["textbox", "shape", "txbx"]):
                        container_id = id(ancestor)
                        break

                if container_id not in container_paragraphs:
                    container_paragraphs[container_id] = []
                container_paragraphs[container_id].append(
                    (element, self._get_paragraph_position(element))
                )

            # Handle txbxContent elements (Word DrawingML textboxes)
            elif tag_name == "txbxContent":
                paragraphs = element.findall(".//w:p", namespaces=element.nsmap)
                container_id = id(element)
                if container_id not in container_paragraphs:
                    container_paragraphs[container_id] = []

                for p in paragraphs:
                    p_id = id(p)
                    if p_id not in processed_paragraphs:
                        processed_paragraphs.append(p_id)
                        container_paragraphs[container_id].append(
                            (p, self._get_paragraph_position(p))
                        )
            else:
                # Try to extract any paragraphs from unknown elements
                paragraphs = element.findall(".//w:p", namespaces=element.nsmap)
                container_id = id(element)
                if container_id not in container_paragraphs:
                    container_paragraphs[container_id] = []

                for p in paragraphs:
                    p_id = id(p)
                    if p_id not in processed_paragraphs:
                        processed_paragraphs.append(p_id)
                        container_paragraphs[container_id].append(
                            (p, self._get_paragraph_position(p))
                        )

        return container_paragraphs

    def _handle_textbox_content(
        self,
        textbox_elements: list,
        doc: DoclingDocument,
    ) -> list[RefItem]:
        elem_ref: list[RefItem] = []
        """Process textbox content and add it to the document structure."""
        level = self._get_level()
        # Create a textbox group to contain all text from the textbox
        textbox_group = doc.add_group(
            label=GroupLabel.SECTION,
            parent=self.parents[level - 1],
            name="textbox",
            content_layer=self.content_layer,
        )
        elem_ref.append(textbox_group.get_ref())
        # Set this as the current parent to ensure textbox content
        # is properly nested in document structure
        original_parent = self.parents[level]
        self.parents[level] = textbox_group

        # Collect and organize paragraphs
        container_paragraphs = self._collect_textbox_paragraphs(textbox_elements)

        # Process all paragraphs
        all_paragraphs = []

        # Sort paragraphs within each container, then process containers
        for paragraphs in container_paragraphs.values():
            # Sort by vertical position within each container
            sorted_container_paragraphs = sorted(
                paragraphs,
                key=lambda x: (
                    x[1] is None,
                    x[1] if x[1] is not None else float("inf"),
                ),
            )

            # Add the sorted paragraphs to our processing list
            all_paragraphs.extend(sorted_container_paragraphs)

        # Track processed paragraphs to avoid duplicates (same content and position)
        processed_paragraphs = set()

        # Process all the paragraphs
        for p, position in all_paragraphs:
            # Create paragraph object to get text content
            paragraph = Paragraph(p, self.docx_obj)
            text_content = paragraph.text

            # Create a unique identifier based on content and position
            paragraph_id = (text_content, position)

            # Skip if this paragraph (same content and position) was already processed
            if paragraph_id in processed_paragraphs:
                _log.debug(
                    f"Skipping duplicate paragraph: content='{text_content[:50]}...', position={position}"
                )
                continue

            # Mark this paragraph as processed
            processed_paragraphs.add(paragraph_id)

            elem_ref.extend(self._handle_text_elements(p, doc))

        # Restore original parent
        self.parents[level] = original_parent
        return elem_ref

    def _handle_equations_in_text(self, element, text):
        only_texts = []
        only_equations = []
        texts_and_equations = []
        for subt in element.iter():
            tag_name = etree.QName(subt).localname
            if tag_name == "t" and "math" not in subt.tag:
                if isinstance(subt.text, str):
                    only_texts.append(subt.text)
                    texts_and_equations.append(subt.text)
            elif "oMath" in subt.tag and "oMathPara" not in subt.tag:
                latex_equation = str(oMath2Latex(subt)).strip()
                if len(latex_equation) > 0:
                    only_equations.append(
                        self.equation_bookends.format(EQ=latex_equation)
                    )
                    texts_and_equations.append(
                        self.equation_bookends.format(EQ=latex_equation)
                    )

        if len(only_equations) < 1:
            return text, []

        if (
            re.sub(r"\s+", "", "".join(only_texts)).strip()
            != re.sub(r"\s+", "", text).strip()
        ):
            # If we are not able to reconstruct the initial raw text
            # do not try to parse equations and return the original
            return text, []

        # Insert equations into original text
        # This is done to preserve white space structure
        output_text = text[:]
        init_i = 0
        for i_substr, substr in enumerate(texts_and_equations):
            if len(substr) == 0:
                continue

            if substr in output_text[init_i:]:
                init_i += output_text[init_i:].find(substr) + len(substr)
            else:
                if i_substr > 0:
                    output_text = output_text[:init_i] + substr + output_text[init_i:]
                    init_i += len(substr)
                else:
                    output_text = substr + output_text

        return output_text, only_equations

    def _create_or_reuse_parent(
        self,
        *,
        doc: DoclingDocument,
        prev_parent: Optional[NodeItem],
        paragraph_elements: list,
    ) -> Optional[NodeItem]:
        return (
            doc.add_inline_group(parent=prev_parent, content_layer=self.content_layer)
            if len(paragraph_elements) > 1
            else prev_parent
        )

    def _handle_text_elements(
        self,
        element: BaseOxmlElement,
        doc: DoclingDocument,
    ) -> list[RefItem]:
        elem_ref: list[RefItem] = []
        paragraph = Paragraph(element, self.docx_obj)
        paragraph_elements = self._get_paragraph_elements(paragraph)
        text, equations = self._handle_equations_in_text(
            element=element, text=paragraph.text
        )

        if text is None:
            return elem_ref
        text = text.strip()

        # Common styles for bullet and numbered lists.
        # "List Bullet", "List Number", "List Paragraph"
        # Identify whether list is a numbered list or not
        p_style_id, p_level = self._get_label_and_level(paragraph)
        numid, ilevel = self._get_numId_and_ilvl(paragraph)

        if numid == 0:
            numid = None

        # Handle lists
        if (
            numid is not None
            and ilevel is not None
            and p_style_id not in ["Title", "Heading"]
        ):
            # Check if this is actually a numbered list by examining the numFmt
            is_numbered = self._is_numbered_list(numid, ilevel)

            li = self._add_list_item(
                doc=doc,
                numid=numid,
                ilevel=ilevel,
                elements=paragraph_elements,
                is_numbered=is_numbered,
            )
            elem_ref.extend(li)  # MUST BE REF!!!
            self._update_history(p_style_id, p_level, numid, ilevel)
            return elem_ref
        elif (
            numid is None
            and self._prev_numid() is not None
            and p_style_id not in ["Title", "Heading"]
        ):  # Close list
            if self.level_at_new_list:
                for key in range(len(self.parents)):
                    if key >= self.level_at_new_list:
                        self.parents[key] = None
                self.level = self.level_at_new_list - 1
                self.level_at_new_list = None
            else:
                for key in range(len(self.parents)):
                    self.parents[key] = None
                self.level = 0

        if p_style_id in ["Title"]:
            for key in range(len(self.parents)):
                self.parents[key] = None
            te = doc.add_text(
                parent=None,
                label=DocItemLabel.TITLE,
                text=text,
                content_layer=self.content_layer,
            )
            self.parents[0] = te
            elem_ref.append(te.get_ref())
        elif "Heading" in p_style_id:
            style_element = getattr(paragraph.style, "element", None)
            if style_element is not None:
                is_numbered_style = (
                    "<w:numPr>" in style_element.xml or "<w:numPr>" in element.xml
                )
            else:
                is_numbered_style = False
            h1 = self._add_heading(doc, p_level, text, is_numbered_style)
            elem_ref.extend(h1)

        elif len(equations) > 0:
            if (paragraph.text is None or len(paragraph.text.strip()) == 0) and len(
                text
            ) > 0:
                # Standalone equation
                level = self._get_level()
                t1 = doc.add_text(
                    label=DocItemLabel.FORMULA,
                    parent=self.parents[level - 1],
                    text=text.replace("<eq>", "").replace("</eq>", ""),
                    content_layer=self.content_layer,
                )
                elem_ref.append(t1.get_ref())
            else:
                # Inline equation
                level = self._get_level()
                inline_equation = doc.add_inline_group(
                    parent=self.parents[level - 1], content_layer=self.content_layer
                )
                elem_ref.append(inline_equation.get_ref())
                text_tmp = text
                for eq in equations:
                    if len(text_tmp) == 0:
                        break

                    split_text_tmp = text_tmp.split(eq.strip(), maxsplit=1)

                    pre_eq_text = split_text_tmp[0]
                    text_tmp = "" if len(split_text_tmp) == 1 else split_text_tmp[1]

                    if len(pre_eq_text) > 0:
                        e1 = doc.add_text(
                            label=DocItemLabel.TEXT,
                            parent=inline_equation,
                            text=pre_eq_text,
                            content_layer=self.content_layer,
                        )
                        elem_ref.append(e1.get_ref())
                    e2 = doc.add_text(
                        label=DocItemLabel.FORMULA,
                        parent=inline_equation,
                        text=eq.replace("<eq>", "").replace("</eq>", ""),
                        content_layer=self.content_layer,
                    )
                    elem_ref.append(e2.get_ref())

                if len(text_tmp) > 0:
                    e3 = doc.add_text(
                        label=DocItemLabel.TEXT,
                        parent=inline_equation,
                        text=text_tmp.strip(),
                        content_layer=self.content_layer,
                    )
                    elem_ref.append(e3.get_ref())

        elif p_style_id in [
            "Paragraph",
            "Normal",
            "Subtitle",
            "Author",
            "DefaultText",
            "ListParagraph",
            "ListBullet",
            "Quote",
        ]:
            level = self._get_level()
            parent = self._create_or_reuse_parent(
                doc=doc,
                prev_parent=self.parents.get(level - 1),
                paragraph_elements=paragraph_elements,
            )
            for text, format, hyperlink in paragraph_elements:
                t2 = doc.add_text(
                    label=DocItemLabel.TEXT,
                    parent=parent,
                    text=text,
                    formatting=format,
                    hyperlink=hyperlink,
                    content_layer=self.content_layer,
                )
                elem_ref.append(t2.get_ref())

        else:
            # Text style names can, and will have, not only default values but user values too
            # hence we treat all other labels as pure text
            level = self._get_level()
            parent = self._create_or_reuse_parent(
                doc=doc,
                prev_parent=self.parents.get(level - 1),
                paragraph_elements=paragraph_elements,
            )
            for text, format, hyperlink in paragraph_elements:
                t3 = doc.add_text(
                    label=DocItemLabel.TEXT,
                    parent=parent,
                    text=text,
                    formatting=format,
                    hyperlink=hyperlink,
                    content_layer=self.content_layer,
                )
                elem_ref.append(t3.get_ref())

        self._update_history(p_style_id, p_level, numid, ilevel)
        return elem_ref

    def _add_heading(
        self,
        doc: DoclingDocument,
        curr_level: Optional[int],
        text: str,
        is_numbered_style: bool = False,
    ) -> list[RefItem]:
        elem_ref: list[RefItem] = []
        level = self._get_level()
        if isinstance(curr_level, int):
            if curr_level > level:
                # add invisible group
                for i in range(level, curr_level):
                    gr1 = doc.add_group(
                        parent=self.parents[i - 1],
                        label=GroupLabel.SECTION,
                        name=f"header-{i}",
                    )
                    elem_ref.append(gr1.get_ref())
                    self.parents[i] = gr1

            elif curr_level < level:
                # remove the tail
                for key in range(len(self.parents)):
                    if key >= curr_level:
                        self.parents[key] = None

            current_level = curr_level
            parent_level = curr_level - 1
            add_level = curr_level
        else:
            current_level = self.level
            parent_level = self.level - 1
            add_level = 1

        if is_numbered_style:
            if add_level in self.numbered_headers:
                self.numbered_headers[add_level] += 1
            else:
                self.numbered_headers[add_level] = 1
            text = f"{self.numbered_headers[add_level]} {text}"

            # Reset deeper levels
            next_level = add_level + 1
            while next_level in self.numbered_headers:
                self.numbered_headers[next_level] = 0
                next_level += 1

            # Scan upper levels
            previous_level = add_level - 1
            while previous_level in self.numbered_headers:
                # MSWord convention: no empty sublevels
                # I.e., sub-sub section (2.0.1) without a sub-section (2.1)
                # is processed as 2.1.1
                if self.numbered_headers[previous_level] == 0:
                    self.numbered_headers[previous_level] += 1

                text = f"{self.numbered_headers[previous_level]}.{text}"
                previous_level -= 1

        hd = doc.add_heading(
            parent=self.parents[parent_level],
            text=text,
            level=add_level,
        )
        self.parents[current_level] = hd
        elem_ref.append(hd.get_ref())
        return elem_ref

    def _add_formatted_list_item(
        self,
        doc: DoclingDocument,
        elements: list,
        marker: str,
        enumerated: bool,
        level: int,
    ) -> list[RefItem]:
        elem_ref: list[RefItem] = []
        # This should not happen by construction
        if not isinstance(self.parents[level], ListGroup):
            _log.warning(
                "Parent element of the list item is not a ListGroup. The list item will be ignored."
            )
            return elem_ref
        if not elements:
            return elem_ref

        if len(elements) == 1:
            text, format, hyperlink = elements[0]
            if text:
                doc.add_list_item(
                    marker=marker,
                    enumerated=enumerated,
                    parent=self.parents[level],
                    text=text,
                    formatting=format,
                    hyperlink=hyperlink,
                )
        else:
            new_item = doc.add_list_item(
                marker=marker,
                enumerated=enumerated,
                parent=self.parents[level],
                text="",
            )
            new_parent = doc.add_inline_group(parent=new_item)
            for text, format, hyperlink in elements:
                if text:
                    doc.add_text(
                        label=DocItemLabel.TEXT,
                        parent=new_parent,
                        text=text,
                        formatting=format,
                        hyperlink=hyperlink,
                        content_layer=self.content_layer,
                    )
        return elem_ref

    def _add_list_item(
        self,
        *,
        doc: DoclingDocument,
        numid: int,
        ilevel: int,
        elements: list,
        is_numbered: bool = False,
    ) -> list[RefItem]:
        elem_ref: list[RefItem] = []
        # this method is always called with is_numbered. Numbered lists should be properly addressed.
        if not elements:
            return elem_ref
        enum_marker = ""

        level = self._get_level()
        prev_indent = self._prev_indent()
        if self._prev_numid() is None or (
            self._prev_numid() == numid and self.level_at_new_list is None
        ):  # Open new list
            self.level_at_new_list = level

            # Reset counters for the new numbering sequence
            self._reset_list_counters_for_new_sequence(numid)

            list_gr = doc.add_list_group(
                name="list",
                parent=self.parents[level - 1],
                content_layer=self.content_layer,
            )
            self.parents[level] = list_gr
            elem_ref.append(list_gr.get_ref())

            # Set marker and enumerated arguments if this is an enumeration element.
            if is_numbered:
                counter = self._get_list_counter(numid, ilevel)
                enum_marker = str(counter) + "."
            else:
                enum_marker = ""
            self._add_formatted_list_item(
                doc, elements, enum_marker, is_numbered, level
            )
        elif (
            self._prev_numid() == numid
            and self.level_at_new_list is not None
            and prev_indent is not None
            and prev_indent < ilevel
        ):  # Open indented list
            for i in range(
                self.level_at_new_list + prev_indent + 1,
                self.level_at_new_list + ilevel + 1,
            ):
                list_gr1 = doc.add_list_group(
                    name="list",
                    parent=self.parents[i - 1],
                    content_layer=self.content_layer,
                )
                self.parents[i] = list_gr1
                elem_ref.append(list_gr1.get_ref())

            # TODO: Set marker and enumerated arguments if this is an enumeration element.
            if is_numbered:
                counter = self._get_list_counter(numid, ilevel)
                enum_marker = str(counter) + "."
            else:
                enum_marker = ""
            self._add_formatted_list_item(
                doc,
                elements,
                enum_marker,
                is_numbered,
                self.level_at_new_list + ilevel,
            )
        elif (
            self._prev_numid() == numid
            and self.level_at_new_list is not None
            and prev_indent is not None
            and ilevel < prev_indent
        ):  # Close list
            for k in self.parents:
                if k > self.level_at_new_list + ilevel:
                    self.parents[k] = None

            # TODO: Set marker and enumerated arguments if this is an enumeration element.
            if is_numbered:
                counter = self._get_list_counter(numid, ilevel)
                enum_marker = str(counter) + "."
            else:
                enum_marker = ""
            self._add_formatted_list_item(
                doc,
                elements,
                enum_marker,
                is_numbered,
                self.level_at_new_list + ilevel,
            )

        elif self._prev_numid() == numid or prev_indent == ilevel:
            # Set marker and enumerated arguments if this is an enumeration element.
            if is_numbered:
                counter = self._get_list_counter(numid, ilevel)
                enum_marker = str(counter) + "."
            else:
                enum_marker = ""
            self._add_formatted_list_item(
                doc, elements, enum_marker, is_numbered, level - 1
            )
        else:
            _log.warning("List item not matching any insert condition.")
        return elem_ref

    @staticmethod
    def _group_cell_elements(
        group_name: str,
        doc: DoclingDocument,
        provs_in_cell: list[RefItem],
        docling_table: TableItem,
        content_layer: ContentLayer = ContentLayer.BODY,
    ) -> RefItem:
        group_element = doc.add_group(
            label=GroupLabel.UNSPECIFIED,
            name=group_name,
            parent=docling_table,
            content_layer=content_layer,
        )
        for prov in provs_in_cell:
            group_element.children.append(prov)
            pr_item = prov.resolve(doc)
            item_parent = pr_item.parent.resolve(doc)
            if pr_item.get_ref() in item_parent.children:
                item_parent.children.remove(pr_item.get_ref())
            pr_item.parent = group_element.get_ref()
        ref_for_rich_cell = group_element.get_ref()
        return ref_for_rich_cell

    def _handle_tables(
        self,
        element: BaseOxmlElement,
        doc: DoclingDocument,
    ) -> list[RefItem]:
        elem_ref: list[RefItem] = []
        table: Table = Table(element, self.docx_obj)
        num_rows = len(table.rows)
        num_cols = len(table.columns)
        _log.debug(f"Table grid with {num_rows} rows and {num_cols} columns")

        if num_rows == 1 and num_cols == 1:
            cell_element = table.rows[0].cells[0]
            # In case we have a table of only 1 cell, we consider it furniture
            # And proceed processing the content of the cell as though it's in the document body
            self._walk_linear(cell_element._element, doc)
            return elem_ref

        data = TableData(num_rows=num_rows, num_cols=num_cols)
        level = self._get_level()
        docling_table = doc.add_table(
            data=data, parent=self.parents[level - 1], content_layer=self.content_layer
        )
        elem_ref.append(docling_table.get_ref())

        cell_set: set[CT_Tc] = set()
        for row_idx, row in enumerate(table.rows):
            _log.debug(f"Row index {row_idx} with {len(row.cells)} populated cells")
            col_idx = 0
            while col_idx < num_cols:
                cell: _Cell = row.cells[col_idx]
                _log.debug(
                    f" col {col_idx} grid_span {cell.grid_span} grid_cols_before {row.grid_cols_before}"
                )
                if cell is None or cell._tc in cell_set:
                    _log.debug("  skipped since repeated content")
                    col_idx += cell.grid_span
                    continue
                else:
                    cell_set.add(cell._tc)

                spanned_idx = row_idx
                spanned_tc: Optional[CT_Tc] = cell._tc
                while spanned_tc == cell._tc:
                    spanned_idx += 1
                    spanned_tc = (
                        table.rows[spanned_idx].cells[col_idx]._tc
                        if spanned_idx < num_rows
                        else None
                    )
                _log.debug(f"  spanned before row {spanned_idx}")

                # Detect equations in cell text
                text, equations = self._handle_equations_in_text(
                    element=cell._element, text=cell.text
                )
                if len(equations) == 0:
                    text = cell.text
                else:
                    text = text.replace("<eq>", "$").replace("</eq>", "$")

                provs_in_cell: list[RefItem] = []
                rich_table_cell: bool = self._is_rich_table_cell(cell)

                if rich_table_cell:
                    _, provs_in_cell = self._walk_linear(cell._element, doc)
                _log.debug(f"Table cell {row_idx},{col_idx} rich? {rich_table_cell}")

                if len(provs_in_cell) > 0:
                    # Cell has multiple elements, we need to group them
                    rich_table_cell = True
                    group_name = f"rich_cell_group_{len(doc.tables)}_{col_idx}_{row.grid_cols_before + row_idx}"
                    ref_for_rich_cell = MsWordDocumentBackend._group_cell_elements(
                        group_name,
                        doc,
                        provs_in_cell,
                        docling_table,
                        content_layer=self.content_layer,
                    )

                if rich_table_cell:
                    rich_cell = RichTableCell(
                        text=text,
                        row_span=spanned_idx - row_idx,
                        col_span=cell.grid_span,
                        start_row_offset_idx=row.grid_cols_before + row_idx,
                        end_row_offset_idx=row.grid_cols_before + spanned_idx,
                        start_col_offset_idx=col_idx,
                        end_col_offset_idx=col_idx + cell.grid_span,
                        column_header=row.grid_cols_before + row_idx == 0,
                        row_header=False,
                        ref=ref_for_rich_cell,  # points to an artificial group around children
                    )
                    doc.add_table_cell(table_item=docling_table, cell=rich_cell)
                    col_idx += cell.grid_span
                else:
                    simple_cell = TableCell(
                        text=text,
                        row_span=spanned_idx - row_idx,
                        col_span=cell.grid_span,
                        start_row_offset_idx=row.grid_cols_before + row_idx,
                        end_row_offset_idx=row.grid_cols_before + spanned_idx,
                        start_col_offset_idx=col_idx,
                        end_col_offset_idx=col_idx + cell.grid_span,
                        column_header=row.grid_cols_before + row_idx == 0,
                        row_header=False,
                    )
                    doc.add_table_cell(table_item=docling_table, cell=simple_cell)
                    col_idx += cell.grid_span
        return elem_ref

    def _has_blip(self, element: BaseOxmlElement) -> bool:
        """Check if a docx element holds any BLIP as a child.

        Args:
            element: a docx element

        Returns:
            Whether the element contains a BLIP as a direct child.
        """

        for item in element:
            if self.blip_xpath_expr(item):
                return True
            if item.findall(
                ".//w:drawing", namespaces=MsWordDocumentBackend._BLIP_NAMESPACES
            ):
                return True

        return False

    def _is_rich_table_cell(self, cell: _Cell) -> bool:
        """Determine whether a docx cell should be parsed as a Docling RichTableCell.

        A docx cell can hold rich content and be parsed with a Docling RichTableCell.
        However, this requires walking through the lxml elements and creating
        node items. If the cell holds only plain text, a TableCell, the parsing
        is simpler and using a TableCell is prefered.

        Plain text means:
        - The cell has only one paragraph
        - The paragraph consists solely of runs with no run properties
          (no need of Docling formatting).
        - No other block-level elements are present inside the cell element.

        Args:
            cell: A docx cell

        Returns:
            Whether the docx cell should be parsed as RichTableCell
        """
        tc = cell._tc

        # must contain only one paragraph
        paragraphs = list(
            tc.iterchildren(
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p"
            )
        )
        if len(paragraphs) > 1:
            return True

        # no other content
        allowed_tags = {"p", "tcPr"}  # paragraph or table-cell properties
        for child in tc:
            tag = child.tag.split("}")[-1]
            if tag not in allowed_tags:
                return True
        if self._has_blip(tc):
            return True

        # paragraph must contain runs with no run-properties
        for para in paragraphs:
            runs = list(
                para.iterchildren(
                    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r"
                )
            )
            for rn in runs:
                item: Run = Run(rn, self.docx_obj)
                if item is not None:
                    fm = MsWordDocumentBackend._get_format_from_run(item)
                    if fm != Formatting():
                        return True

        # All checks passed: plain text only
        return False

    def _handle_pictures(
        self, drawing_blip: Any, doc: DoclingDocument
    ) -> list[RefItem]:
        def get_docx_image(drawing_blip: Any) -> Optional[bytes]:
            image_data: Optional[bytes] = None
            rId = drawing_blip[0].get(
                "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
            )
            if rId in self.docx_obj.part.rels:
                # Access the image part using the relationship ID
                image_part = self.docx_obj.part.rels[rId].target_part
                image_data = image_part.blob  # Get the binary image data
            return image_data

        elem_ref: list[RefItem] = []
        level = self._get_level()
        # Open the BytesIO object with PIL to create an Image
        image_data: Optional[bytes] = get_docx_image(drawing_blip)
        if image_data is None:
            _log.warning("Warning: image cannot be found")
            p1 = doc.add_picture(
                parent=self.parents[level - 1],
                caption=None,
                content_layer=self.content_layer,
            )
            elem_ref.append(p1.get_ref())
        else:
            try:
                image_bytes = BytesIO(image_data)
                pil_image = Image.open(image_bytes)
                p2 = doc.add_picture(
                    parent=self.parents[level - 1],
                    image=ImageRef.from_pil(image=pil_image, dpi=72),
                    caption=None,
                    content_layer=self.content_layer,
                )
                elem_ref.append(p2.get_ref())
            except (UnidentifiedImageError, OSError):
                _log.warning("Warning: image cannot be loaded by Pillow")
                p3 = doc.add_picture(
                    parent=self.parents[level - 1],
                    caption=None,
                    content_layer=self.content_layer,
                )
                elem_ref.append(p3.get_ref())
        return elem_ref

    def _handle_drawingml(self, doc: DoclingDocument, drawingml_els: Any):
        # 1) Make an empty copy of the original document
        dml_doc = self.load_msword_file(self.path_or_stream, self.document_hash)
        body = dml_doc._element.body
        for child in list(body):
            body.remove(child)

        # 2) Add DrawingML to empty document
        new_para = dml_doc.add_paragraph()
        new_r = new_para.add_run()
        for dml in drawingml_els:
            new_r._r.append(deepcopy(dml))

        # 3) Export DOCX->PDF->PNG and save it in DoclingDocument
        level = self._get_level()
        try:
            pil_image = get_pil_from_dml_docx(
                dml_doc, converter=self.docx_to_pdf_converter
            )
            if pil_image is None:
                raise UnidentifiedImageError

            doc.add_picture(
                parent=self.parents[level - 1],
                image=ImageRef.from_pil(image=pil_image, dpi=72),
                caption=None,
                content_layer=self.content_layer,
            )
        except (UnidentifiedImageError, OSError):
            _log.warning("Warning: DrawingML image cannot be loaded by Pillow")
            doc.add_picture(
                parent=self.parents[level - 1],
                caption=None,
                content_layer=self.content_layer,
            )

        return

    def _add_header_footer(self, docx_obj: DocxDocument, doc: DoclingDocument) -> None:
        """Add section headers and footers.

        Headers and footers are added in the furniture content and only the text paragraphs
        are parsed. The paragraphs are attached to a single group item for the header or the
        footer. If the document has a section with new header and footer, they will be parsed
        in new group items.

        Args:
            docx_obj: A docx Document object to be parsed.
            doc: A DoclingDocument object to add the header and footer from docx_obj.
        """
        current_layer = self.content_layer
        base_parent = self.parents[0]
        self.content_layer = ContentLayer.FURNITURE
        for sec_idx, section in enumerate(docx_obj.sections):
            if sec_idx > 0 and not section.different_first_page_header_footer:
                continue

            hdr = (
                section.first_page_header
                if section.different_first_page_header_footer
                else section.header
            )
            par = [txt for txt in (par.text.strip() for par in hdr.paragraphs) if txt]
            tables = hdr.tables
            has_blip = self._has_blip(hdr._element)
            if par or tables or has_blip:
                self.parents[0] = doc.add_group(
                    label=GroupLabel.SECTION,
                    name="page header",
                    content_layer=self.content_layer,
                )
                self._walk_linear(hdr._element, doc)

            ftr = (
                section.first_page_footer
                if section.different_first_page_header_footer
                else section.footer
            )
            par = [txt for txt in (par.text.strip() for par in ftr.paragraphs) if txt]
            tables = ftr.tables
            has_blip = self._has_blip(ftr._element)
            if par or tables or has_blip:
                self.parents[0] = doc.add_group(
                    label=GroupLabel.SECTION,
                    name="page footer",
                    content_layer=self.content_layer,
                )
                self._walk_linear(ftr._element, doc)

        self.content_layer = current_layer
        self.parents[0] = base_parent
